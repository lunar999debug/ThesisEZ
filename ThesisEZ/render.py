"""学位论文渲染引擎 —— 基于 pandoc 的纯 Python 渲染脚本。

将 index.md + 各章节 .md（按文件名排序）合并成单一源文件，使用 Word
模板渲染为 docx，再做后处理：表格样式、题注、图片环绕方式、按章
节插入分节符等。模板可替换为任意学校 / 任意格式的 .docx 模板。

用法：
    python render.py                       # docx (zh, 默认)，CSL=numeric
    python render.py docx
    python render.py docx-en
    python render.py --csl=note            # numeric (默认) | note | author-date
    python render.py docx author-date      # 也接受裸 CSL 名称
    python render.py --no-bibtex           # ★ 跳过 BibTeX/citeproc，参考文献按纯文本走
"""
import copy
import datetime
import glob
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml


# 后处理过程中需要用到的 XML 命名空间 URI。
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DRAW_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


# ── 通用 XML 工具函数 ───────────────────────────────────

def _build_style_tables(doc):
    """根据 doc.styles 构造 (样式名→样式ID, 一级标题样式ID集合)。"""
    name_to_id = {s.name: s.style_id for s in doc.styles}
    h1_style_ids = {
        name_to_id.get("Heading 1", "Heading1"),
        "Heading1", "Heading 1", "1",
        name_to_id.get("非编号章节标题"),
        name_to_id.get("非编号章节标题（目录不显示）"),
    }
    h1_style_ids.discard(None)
    return name_to_id, h1_style_ids


def _ensure_pPr(p):
    """返回 p 的 <w:pPr> 子节点，若不存在则创建。"""
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    return pPr


def _assign_para_style(p, style_id: str) -> None:
    """设置段落样式，覆盖原有的 pStyle。"""
    pPr = _ensure_pPr(p)
    for legacy in pPr.findall(qn("w:pStyle")):
        pPr.remove(legacy)
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)


def _read_para_style(p):
    """返回段落的样式 ID，找不到则返回 None。"""
    ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else None


def _gather_para_text(p) -> str:
    """把 p 中所有 <w:t> 的文本拼成一个字符串返回。"""
    return "".join((t.text or "") for t in p.findall(".//" + qn("w:t")))


def _push_text_run(parent, text: str, preserve: bool = False) -> None:
    """向 parent 追加一个 <w:r><w:t>text</w:t></w:r>。"""
    r = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    if preserve:
        t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    r.append(t_el)
    parent.append(r)


def _new_page_break_para():
    """构造一个只含 <w:br w:type='page'/> 的空段落。"""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _emit_seq_field(parent, seq_name: str) -> None:
    """向 parent 追加 SEQ 域的 begin/instrText/end 三段 run。"""
    f1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    f1.append(fc1)
    parent.append(f1)

    f2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    f2.append(instr)
    parent.append(f2)

    f3 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    f3.append(fc2)
    parent.append(f3)


# 手动编号前缀 —— 与 aichuli.py 中保持同步。如果题注文本以
# "@@NUM@@N-M|description" 开头，表示使用手动编号，不插入 SEQ 项。
MANUAL_NUM_PREFIX = "@@NUM@@"


def _parse_manual_num(text: str):
    """如果 text 以 MANUAL_NUM_PREFIX 开头，返回 (编号, 描述)；
    否则返回 None。"""
    if not text or not text.startswith(MANUAL_NUM_PREFIX):
        return None
    rest = text[len(MANUAL_NUM_PREFIX):]
    if "|" in rest:
        num, desc = rest.split("|", 1)
    else:
        num, desc = rest, ""
    return num.strip(), desc.strip()


def _build_caption_para(label: str, seq_name: str, text: str,
                        caption_style_id: str):
    """构造题注段落。两种模式：
    1. 常规模式：'<label> { SEQ <seq_name> } <text>'，自动递增编号。
    2. 手动编号模式：当 text 以 @@NUM@@ 开头时，输出为
       '<label> <手写编号> <描述>'，例如 '图 3-1 实验流程图'。
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), caption_style_id)
    pPr.append(pStyle)
    p.append(pPr)

    parsed = _parse_manual_num(text)
    if parsed is not None:
        num, desc = parsed
        body = f"{label} {num}"
        if desc:
            body += f"　{desc}"  # 全角空格分隔，符合中文题注习惯
        _push_text_run(p, body, preserve=True)
        return p

    _push_text_run(p, f"{label} ", preserve=True)
    _emit_seq_field(p, seq_name)
    if text:
        _push_text_run(p, f" {text}", preserve=True)
    return p


def _rewrite_sectpr_rids(sectpr, tpl_rid_to_target, out_target_to_rid):
    """重映射克隆 sectPr 中的 header/footer rId；找不到映射的引用直接丢弃。"""
    for ref in list(sectpr):
        if ref.tag not in (qn("w:headerReference"), qn("w:footerReference")):
            continue
        tpl_rid = ref.get(f"{{{_REL_NS}}}id")
        target = tpl_rid_to_target.get(tpl_rid)
        out_rid = out_target_to_rid.get(target) if target else None
        if out_rid:
            ref.set(f"{{{_REL_NS}}}id", out_rid)
        else:
            sectpr.remove(ref)


def locate_pandoc() -> str:
    """定位 pandoc 可执行文件路径。优先查找 conda 环境下的 envs/pandoc/。"""
    py_dir = Path(sys.executable).parent
    # 向上回溯找 conda 根目录（含有 envs/ 子目录），优先返回
    # `<root>/envs/pandoc/Library/bin/pandoc.exe`。无论当前解释器是
    # base 环境还是子环境（如 envs/mybase）都能正确识别。
    preferred = []
    for anc in [py_dir, *py_dir.parents]:
        cand = anc / "envs" / "pandoc" / "Library" / "bin" / "pandoc.exe"
        if cand.is_file():
            preferred.append(cand)
            break
    preferred += [Path(p) for p in glob.glob(
        str(py_dir / "envs" / "*" / "Library" / "bin" / "pandoc.exe")
    )]
    for c in preferred:
        if c.is_file():
            return str(c)
    exe = shutil.which("pandoc")
    if exe:
        return exe
    candidates = [
        py_dir / "pandoc.exe",
        py_dir / "Scripts" / "pandoc.exe",
        py_dir / "Library" / "bin" / "pandoc.exe",
    ]
    candidates += [Path(p) for p in glob.glob(
        str(py_dir / "envs" / "*" / "Scripts" / "pandoc.exe")
    )]
    for c in candidates:
        if c.is_file():
            return str(c)
    sys.exit("未找到 pandoc。")


PROJECT_ROOT = Path(__file__).resolve().parent
MD_DIR = PROJECT_ROOT / "md"
BIB = PROJECT_ROOT / "401-reference.bib"
OUT_BASE = PROJECT_ROOT / "ThesisEZ"

# 合并源文件时使用的 YAML 头信息。index.md 已不再读取——
# 标题/作者来自模板§1（详见 _inject_template_section1），此处只保留与
# citeproc 相关的元数据。
FRONT_MATTER_BIB = (
    "---\n"
    f"bibliography: [{BIB.name}]\n"
    "nocite: '@*'\n"
    "link-citations: yes\n"
    "---\n"
)

# --no-bibtex 模式下使用的最简化 front-matter（不引入 bib 引擎）。
FRONT_MATTER_PLAIN = (
    "---\n"
    "link-citations: yes\n"
    "---\n"
)


def _label_top_headings(text: str, style: str) -> str:
    # 给每个一级 ATX 标题加上 docx 自定义样式 + .unnumbered 属性。
    def repl(m: "re.Match[str]") -> str:
        title = m.group(1).strip()
        # 已经带属性块的不要重复打标。
        if title.endswith("}"):
            return m.group(0)
        return f'# {title} {{.unnumbered custom-style="{style}"}}'
    return re.sub(r"^#\s+(.+?)\s*$", repl, text, flags=re.MULTILINE)


def fuse_sources(docx_mode: bool = False, use_bibtex: bool = True) -> Path:
    """合并 md/ 下的所有 markdown 文件为一个临时源文件。"""
    front = FRONT_MATTER_BIB if use_bibtex else FRONT_MATTER_PLAIN
    parts = [front]
    for f in sorted(MD_DIR.glob("*.md")):
        body = f.read_text(encoding="utf-8")
        if docx_mode:
            stem = f.stem
            prefix = stem.split("-", 1)[0] if "-" in stem else stem
            # 文件名前缀里第一个非零数字 → 1-based 模板节号。
            # 每节的一级标题样式在后处理阶段才应用（pandoc 会忽略
            # 标题上的 custom-style）。
            digit = next((c for c in prefix if c.isdigit() and c != "0"), "1")
            # 插入 md-boundary 标记，让后处理阶段知道每章应该归属哪个
            # 模板节，从而继承对应节的页眉页脚。
            marker = (
                f'\n\n::: {{custom-style="RmdMarker{digit}"}}\n'
                f"§\n"
                f":::\n\n"
            )
            body = marker + body
        parts.append(body)
    merged = "\n\n".join(parts)
    # 剥离遗留代码块（```{r ...} ... ```），保留行内文本。
    merged = re.sub(
        r"^```\{[^}]*\}.*?^```\s*$", "", merged,
        flags=re.MULTILINE | re.DOTALL,
    )
    # 把行内 `r format(Sys.Date(), format='...')` 替换成今天日期。
    today = datetime.date.today()
    def _eval_inline_r(m: "re.Match[str]") -> str:
        body = m.group(1)
        fmt_match = re.search(
            r"format\s*\(\s*Sys\.Date\(\)\s*,\s*format\s*=\s*['\"]([^'\"]+)['\"]",
            body,
        )
        if fmt_match:
            # R 的 strftime 格式占位符与 Python 完全一致。
            return today.strftime(fmt_match.group(1))
        if "Sys.Date()" in body:
            return today.isoformat()
        return ""
    merged = re.sub(r"`r\s+([^`]*)`", _eval_inline_r, merged)
    tmp = PROJECT_ROOT / "_merged.md"
    tmp.write_text(merged, encoding="utf-8")
    return tmp


CSL_OPTIONS = {
    "note": PROJECT_ROOT / "china-national-standard-gb-t-7714-2015-note.csl",
    "numeric": PROJECT_ROOT / "china-national-standard-gb-t-7714-2015-numeric.csl",
    "author-date": PROJECT_ROOT / "china-national-standard-gb-t-7714-2015-author-date.csl",
}

def render_docx(pandoc: str, src: Path, lang: str = "zh", csl: str = "numeric",
                use_bibtex: bool = True) -> None:
    """调用 pandoc 把合并后的 md 渲染成 docx，再做后处理。"""
    ref = PROJECT_ROOT / "template" / (
        "中文毕业设计模板260501.docx" if lang == "zh"
        else "英文毕业设计模板250928.docx"
    )
    out = OUT_BASE.with_suffix(".docx") if lang == "zh" else OUT_BASE.with_name(OUT_BASE.name + "_en").with_suffix(".docx")
    cmd = [
        pandoc, str(src), "-o", str(out),
        "--top-level-division=chapter",
        "--toc", "--toc-depth=3",
        f"--reference-doc={ref}",
        "--standalone",
    ]
    if use_bibtex:
        csl_path = CSL_OPTIONS.get(csl)
        if csl_path is None:
            sys.exit(f"未知的 CSL 名称 '{csl}'。可选：{', '.join(CSL_OPTIONS)}")
        cmd[5:5] = [
            f"--bibliography={BIB}",
            f"--csl={csl_path}",
            "--citeproc",
        ]
    print("执行命令:", " ".join(cmd))
    subprocess.run(cmd, cwd=str(PROJECT_ROOT), check=True)
    if lang == "zh":
        _postprocess_docx(out)


def _postprocess_docx(path: Path) -> None:
    """对渲染后的 docx 做一系列后处理：
    应用三线表 + 表格样式，在表格上方/图片下方插入题注样式段落
    （带自动递增的 SEQ 域）等等。
    """

    doc = DocxDocument(str(path))
    name_to_id, h1_style_ids = _build_style_tables(doc)
    style_names = set(name_to_id.keys())

    TABLE_STYLE = "三线表"
    CELL_PARA_STYLE = "表格"
    CAPTION_STYLE = "Caption"  # 兜底
    for s in doc.styles:
        if s.name in ("Caption", "题注"):
            CAPTION_STYLE = s.style_id
            break

    # 1) 表格：套用三线表表样，单元格段落套用「表格」段落样式。
    if TABLE_STYLE in style_names:
        for tbl in doc.tables:
            try:
                tbl.style = doc.styles[TABLE_STYLE]
            except KeyError:
                pass
            if CELL_PARA_STYLE in style_names:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            try:
                                p.style = doc.styles[CELL_PARA_STYLE]
                            except KeyError:
                                pass

    body = doc.element.body

    # 2) 表格题注：保证每张表上方都有一个题注段落。
    #    Pandoc 会在表格上方生成 "TableCaption" 样式的段落（含原始文字），
    #    这里把它改成 "Caption" 样式并前置 "表 { SEQ Table } "。如果原本
    #    没有题注，则插入一个空题注。
    for tbl in body.findall(qn("w:tbl")):
        prev = tbl.getprevious()
        existing_cap = None
        if prev is not None and prev.tag == qn("w:p"):
            ps = prev.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is not None and ps.get(qn("w:val")) in {
                "TableCaption", "Table Caption", CAPTION_STYLE
            }:
                existing_cap = prev
        if existing_cap is not None:
            # 提取原题注文字，替换为 "表 SEQ <text>" 题注段落。
            text = _gather_para_text(existing_cap).strip()
            existing_cap.addprevious(_build_caption_para("表", "Table", text, CAPTION_STYLE))
            body.remove(existing_cap)
        else:
            tbl.addprevious(_build_caption_para("表", "Table", "", CAPTION_STYLE))

    # 3) 图题注：pandoc 会把行内图片放在 "CaptionedFigure" 样式的段落中，
    #    alt 文本存放在 <wp:docPr @descr> 上。我们在这种段落紧后面插入
    #    一个兄弟节点 "Caption" 段落作为题注。
    for p in list(body.findall(qn("w:p"))):
        if _read_para_style(p) not in {"CaptionedFigure", "Captioned Figure"}:
            continue
        descr = ""
        docpr = p.find(".//{%s}docPr" % _DRAW_NS)
        if docpr is not None:
            descr = (docpr.get("descr") or docpr.get("title") or "").strip()
        p.addnext(_build_caption_para("图", "Figure", descr, CAPTION_STYLE))

    # 4) 移除 pandoc 残留的 "ImageCaption" 段落（已被我们换掉）。
    for p in list(body.findall(qn("w:p"))):
        if _read_para_style(p) in {"ImageCaption", "Image Caption"}:
            body.remove(p)

    # 5) 把所有行内图片转换为「上下型环绕」(Top-and-Bottom) 锚定图片。
    #    wp:anchor 的子元素必须严格按 schema 顺序排列：
    #    simplePos, positionH, positionV, extent, effectExtent, wrap*, docPr,
    #    cNvGraphicFramePr?, graphic。一旦顺序错乱，Word 会拒绝打开文件
    #    （或者悄悄丢掉图片）。
    wp = lambda tag: f"{{{_DRAW_NS}}}{tag}"
    for inline in list(body.findall(f".//{wp('inline')}")):
        # 从 inline 元素中拆出必需/可选的子节点。
        extent = inline.find(wp("extent"))
        effect_extent = inline.find(wp("effectExtent"))
        doc_pr = inline.find(wp("docPr"))
        cnv_gfp = inline.find(wp("cNvGraphicFramePr"))
        graphic = inline.find(f"{{{_DML_NS}}}graphic")
        if extent is None or doc_pr is None or graphic is None:
            continue  # 结构异常，原样保留

        anchor = parse_xml(
            f'<wp:anchor xmlns:wp="{_DRAW_NS}"'
            ' distT="0" distB="0" distL="114300" distR="114300"'
            ' simplePos="0" relativeHeight="251658240" behindDoc="0"'
            ' locked="0" layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="column">'
            '<wp:align>center</wp:align>'
            '</wp:positionH>'
            '<wp:positionV relativeFrom="paragraph">'
            '<wp:posOffset>0</wp:posOffset>'
            '</wp:positionV>'
            '<wp:wrapTopAndBottom/>'
            '</wp:anchor>'
        )
        # 在 wrapTopAndBottom 之前插入 extent/effectExtent（索引 3）。
        wrap_el = anchor.find(wp("wrapTopAndBottom"))
        wrap_idx = list(anchor).index(wrap_el)
        insert_at = wrap_idx
        anchor.insert(insert_at, extent)
        insert_at += 1
        if effect_extent is not None:
            anchor.insert(insert_at, effect_extent)
            insert_at += 1
        # docPr / cNvGraphicFramePr / graphic 放到 wrapTopAndBottom 之后。
        anchor.append(doc_pr)
        if cnv_gfp is not None:
            anchor.append(cnv_gfp)
        anchor.append(graphic)

        inline.getparent().replace(inline, anchor)

    # 6) 参考文献：把「参考文献」标题之后的段落统一改为「列表段落（无编号）」样式。
    bib_target_id = name_to_id.get("列表段落（无编号）")
    if bib_target_id:
        body_ps = body.findall(qn("w:p"))
        ref_idx = None
        for i, p in enumerate(body_ps):
            if _read_para_style(p) not in h1_style_ids:
                continue
            txt = _gather_para_text(p).replace(" ", "").replace("\u3000", "")
            if txt == "参考文献":
                ref_idx = i
                break

        if ref_idx is not None:
            for p in body_ps[ref_idx + 1:]:
                if _read_para_style(p) in h1_style_ids:
                    break  # 进入下一章 —— 停止
                _assign_para_style(p, bib_target_id)

    # 7) 图表注：以「注：」/「Note:」开头的段落 → 应用「图表注」样式。
    #    iter() 会同时遍历正文和表格单元格里的段落，一次性处理完。
    note_style_id = name_to_id.get("图表注")
    if note_style_id:
        _NOTE_PREFIXES = ("注：", "注:", "Note:", "Note：")
        for p in body.iter(qn("w:p")):
            txt = _gather_para_text(p).lstrip()
            if any(txt.startswith(pfx) for pfx in _NOTE_PREFIXES):
                _assign_para_style(p, note_style_id)

    # 8) 把 pandoc 自动生成的目录（一段 Word TOC SDT）移动到摘要和正文
    #    之间（即 RmdMarker3 段之前），并把标题改为「目 录」。这一步必须
    #    在 §1 注入之前完成，否则 SDT 会在第 9 步被一起删掉。
    _relocate_toc(doc)

    # 9) 把模板§1（封面 + 英文封面 + 原创性声明 + 授权声明）原样注入到
    #    输出文档最前面，期间把 101-setup.md 配置的字段替换进去。这会
    #    把 pandoc 在§2 之前生成的所有内容都替换掉。
    template_doc = PROJECT_ROOT / "template" / "中文毕业设计模板260501.docx"
    setup_md = PROJECT_ROOT / "101-setup.md"
    if template_doc.is_file() and setup_md.is_file():
        _inject_template_section1(
            doc, template_doc, _parse_setup_table(setup_md)
        )

    # 10) 按 Rmd 文件分节：每一章继承模板对应节的页眉页脚。RmdMarkerN 标记
    #     用来标识章归属。
    if template_doc.is_file():
        _apply_per_rmd_sections(doc, template_doc)

    # 11) 在每个一级标题前都加一个分页符（每个 # 之前分页）。第一个标题
    #     和已经处于新页起始位置（前一段携带 sectPr 或本身含 page break）
    #     的标题会被跳过。
    _ensure_page_break_before_headings(doc)

    # 12) 公式：给独立公式自动编号（单行公式自动编号）。
    #     一个段落如果含 m:oMath 且没有其他文本，则视为独立公式。
    #     应用「公式」样式：居中-tab + 公式 + 右-tab + (SEQ Equation)。
    for p in body.findall(qn("w:p")):
        if not p.findall(".//" + qn("m:oMath")):
            continue
        if _gather_para_text(p).strip():  # 含非数学文本 → 行内公式，跳过
            continue

        # 模板中存在「公式」样式时套用。
        if "公式" in name_to_id:
            _assign_para_style(p, name_to_id["公式"])

        # 设置 tab stop：居中 ~8cm，右对齐 ~16cm。
        pPr = _ensure_pPr(p)
        for old_tag in ("w:jc", "w:tabs"):
            for old in pPr.findall(qn(old_tag)):
                pPr.remove(old)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
        for val, pos in (("center", "4536"), ("right", "9072")):
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), val)
            tab.set(qn("w:pos"), pos)
            tabs.append(tab)

        # 插入开头的 tab（让公式居中），再追加结尾 tab + "(SEQ)"。
        r_tab_start = OxmlElement("w:r")
        r_tab_start.append(OxmlElement("w:tab"))
        p.insert(1, r_tab_start)  # 紧跟在 pPr 之后

        r_tab_end = OxmlElement("w:r")
        r_tab_end.append(OxmlElement("w:tab"))
        p.append(r_tab_end)

        _push_text_run(p, "(", preserve=True)
        _emit_seq_field(p, "Equation")
        _push_text_run(p, ")", preserve=True)

    # 13) 关键词：以「关键词」/「keywords」开头的段 → 应用「关键词」样式。
    kw_style_id = name_to_id.get("关键词")
    if kw_style_id:
        _KW_PREFIXES = ("关键词：", "关键词:", "keywords:", "keywords：")
        for p in body.findall(qn("w:p")):
            txt = _gather_para_text(p).lstrip().lower()
            if any(txt.startswith(pfx) for pfx in _KW_PREFIXES):
                _assign_para_style(p, kw_style_id)

    # 14) settings.xml：设置 <w:updateFields w:val="true"/>，让 Word 在
    #     首次打开时自动刷新目录 / SEQ / 页码等域（打开时自动更新目录）。
    upd = doc.settings.element.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        doc.settings.element.append(upd)
    upd.set(qn("w:val"), "true")

    doc.save(str(path))
    print(f"后处理完成 {path.name}：已应用三线表 / 表格样式，插入题注，图片改为上下型环绕，按章节分节。")


def _parse_setup_table(path: Path) -> "dict[str, str]":
    """解析 101-setup.md 中的三列 markdown 表 → {old_text: new_text}。
    单元格内容可能跨多行（保留行内换行），所以调用方可以按 '\\n' 切分
    来处理多行键值。
    """
    text = path.read_text(encoding="utf-8")
    rows = re.findall(r"\|([^|]+)\|([^|]+)\|([^|]+)\|", text)
    out: dict[str, str] = {}
    saw_header = False
    for c1, c2, c3 in rows:
        c1s = c1.strip()
        if not saw_header and "需填写字段" in c1s:
            saw_header = True
            continue
        if all(s.strip() and set(s.strip()) <= {"-", ":"} for s in (c1, c2, c3)):
            continue  # markdown 表格分隔行
        # 把单元格内的空白都压缩为一个空格 —— 模板现在已是单字段单行，
        # 任何嵌入换行都只是排版需要。
        old = re.sub(r"\s+", " ", c2).strip()
        new = re.sub(r"\s+", " ", c3).strip()
        if old and new and old != new:
            out[old] = new
    return out


def _inject_template_section1(doc, template_path: Path, replacements) -> None:
    """把模板§1（截至并包含第一个分节段落）拷贝进 doc，替换掉首个
    RmdMarkerN 段之前的全部正文内容。同时根据 `replacements` 做文本
    替换（多行键按 '\\n' 拆分：新值放进第一个匹配段，后续行清空）。
    """
    tpl = DocxDocument(str(template_path))
    tpl_body = tpl.element.body
    tpl_rid_to_target = {
        rid: rel.target_ref for rid, rel in tpl.part.rels.items()
    }
    out_target_to_rid = {
        rel.target_ref: rid for rid, rel in doc.part.rels.items()
    }

    # 1) 收集§1 子节点（深拷贝），直到遇到 pPr 中含第一个 sectPr 的段落
    #    （这表示§1 边界）为止，包含该段。
    sec1: list = []
    found_terminator = False
    for child in tpl_body:
        if child.tag == qn("w:sectPr"):
            break  # body 级 sectPr —— 行内未找到§1 边界
        clone = copy.deepcopy(child)
        sec1.append(clone)
        if child.tag == qn("w:p"):
            sectPr = child.find(qn("w:pPr") + "/" + qn("w:sectPr"))
            if sectPr is not None:
                clone_sectPr = clone.find(qn("w:pPr") + "/" + qn("w:sectPr"))
                _rewrite_sectpr_rids(clone_sectPr, tpl_rid_to_target, out_target_to_rid)
                found_terminator = True
                break
    if not found_terminator:
        return

    # 2) 按段落粒度做文本替换。模板已扁平化为「每个字段单独成段」，所以
    #    不再需要做跨段窗口匹配。每段内会替换所有出现位置（例如
    #    "20XX" 在 "20XX年XX月" 和 "June, 20XX" 两处都会出现）。
    def set_text(p, value: str) -> None:
        first_set = False
        for r in p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                if not first_set:
                    t.text = value
                    t.set(qn("xml:space"), "preserve")
                    first_set = True
                else:
                    t.text = ""
        if not first_set and value:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = value
            r.append(t)
            p.append(r)

    all_paragraphs = [p for el in sec1 for p in el.iter(qn("w:p"))]
    for old, new in replacements.items():
        old_s, new_s = old.strip(), new.strip()
        if not old_s:
            continue
        for p in all_paragraphs:
            pt = _gather_para_text(p)
            if old_s in pt:
                set_text(p, pt.replace(old_s, new_s))

    # 3) 输出文档：先删掉所有早于第一个 RmdMarkerN 段的 body 子节点，
    #    然后把§1 整体前置。
    body = doc.element.body
    first_marker = None
    for c in list(body):
        if c.tag != qn("w:p"):
            continue
        sval = _read_para_style(c)
        if sval and sval.startswith("RmdMarker"):
            first_marker = c
            break
    if first_marker is None:
        return
    for c in list(body):
        if c is first_marker:
            break
        body.remove(c)
    for el in sec1:
        first_marker.addprevious(el)


def _ensure_page_break_before_headings(doc) -> None:
    """在每一个一级标题段落之前都插入一个空白页面分隔段
    （Heading 1 / 非编号章节标题 / 非编号章节标题（目录不显示））——
    无条件插入，包括首个标题，也包括前面已有分节符或分页符的标题。
    幂等性：若紧邻的前一段就是我们之前插入的纯分页符段落，则跳过。
    """
    body = doc.element.body
    name_to_id, heading_ids = _build_style_tables(doc)
    # 同样把 ABSTRACT 样式纳入。
    abstract_id = name_to_id.get("ABSTRACT")
    if abstract_id:
        heading_ids.add(abstract_id)

    def is_pure_page_break_para(p) -> bool:
        # 我们插入的段落特征：恰好一个 <w:r>，里面只含一个 <w:br type=page/>。
        runs = p.findall(qn("w:r"))
        if len(runs) != 1:
            return False
        children = list(runs[0])
        if len(children) != 1:
            return False
        c = children[0]
        return c.tag == qn("w:br") and c.get(qn("w:type")) == "page"

    def is_section_break_para(p) -> bool:
        # 含 sectPr 的段落本身就会开启新的一节（也意味着新的一页）。
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            return False
        return pPr.find(qn("w:sectPr")) is not None

    headings = []
    for p in body.findall(qn("w:p")):
        if _read_para_style(p) in heading_ids:
            headings.append(p)

    for p in headings:
        prev = p.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is not None and is_pure_page_break_para(prev):
            continue  # 前面已有我们的分页段落
        if prev is not None and is_section_break_para(prev):
            continue  # 分节符已经会开启新的一页（每节首章不再加分页符）
        p.addprevious(_new_page_break_para())


def _relocate_toc(doc) -> None:
    """把 pandoc 自动生成的目录（一段 Word TOC SDT）移动到摘要与首章
    正文之间。同时把 'Table of Contents' 标题替换为「目录」（套用
    非编号章节标题（目录不显示）样式），并在前面加一个分页段落，让
    目录从新的一页开始。
    """
    body = doc.element.body
    sdt = body.find(qn("w:sdt"))
    if sdt is None:
        return

    # 插入位置：RmdMarker3 段（即「绪论」开始处的标记）。
    target = None
    for p in body.findall(qn("w:p")):
        ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if ps is not None and ps.get(qn("w:val")) == "RmdMarker3":
            target = p
            break
    if target is None:
        return

    # 改写 SDT 内 TOC 标题的样式与文本。
    name_to_id, _ = _build_style_tables(doc)
    hidden_id = name_to_id.get("非编号章节标题（目录不显示）")
    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is not None:
        for p in sdt_content.findall(qn("w:p")):
            ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is None or ps.get(qn("w:val")) != "TOCHeading":
                continue
            if hidden_id:
                ps.set(qn("w:val"), hidden_id)
            # 把第一个文本 run 改成「目 录」，其余清掉。
            runs = p.findall(qn("w:r"))
            for i, r in enumerate(runs):
                if i == 0:
                    for t in r.findall(qn("w:t")):
                        t.text = "目 录"
                        for extra in r.findall(qn("w:t"))[1:]:
                            r.remove(extra)
                        break
                else:
                    p.remove(r)
            break

    page_break_p = _new_page_break_para()

    # 拆出后重新插入：分页符 → SDT → 空锚段 → marker。
    # 末尾留一个空 <w:p> 作为锚点，是为了让 _apply_per_rmd_sections 把
    # §2 的 sectPr 落在 TOC 之后的段落上（否则它会向上回溯越过 SDT
    # —— 一个非 w:p 节点会被跳过 —— 落到 TOC 之前的分页段，从而把
    # 整个 TOC 归到§3 里去）。
    toc_anchor_p = OxmlElement("w:p")
    body.remove(sdt)
    target.addprevious(page_break_p)
    target.addprevious(sdt)
    target.addprevious(toc_anchor_p)


def _apply_per_rmd_sections(doc, template_path: Path) -> None:
    """在章边界处插入分节符，并按对应模板节继承页眉页脚。
    """
    tpl = DocxDocument(str(template_path))
    tpl_sectprs = [s._sectPr for s in tpl.sections]
    tpl_rid_to_target = {
        rid: rel.target_ref for rid, rel in tpl.part.rels.items()
    }
    out_target_to_rid = {
        rel.target_ref: rid for rid, rel in doc.part.rels.items()
    }

    if not tpl_sectprs:
        return

    body = doc.element.body
    name_to_id, _ = _build_style_tables(doc)
    HEADING1_ID = name_to_id.get("Heading 1", "Heading1")
    UNNUMBERED_ID = name_to_id.get("非编号章节标题")
    UNNUMBERED_HIDDEN_ID = name_to_id.get("非编号章节标题（目录不显示）")
    ABSTRACT_ID = name_to_id.get("ABSTRACT")
    HEADING2_ID = name_to_id.get("Heading 2", "Heading2")
    HEADING3_ID = name_to_id.get("Heading 3", "Heading3")
    APPENDIX_H2_ID = name_to_id.get("附录标题2")
    APPENDIX_H3_ID = name_to_id.get("附录标题3")
    HEADING_STYLES = {HEADING1_ID, "Heading1", "Heading 1", "1"}
    if UNNUMBERED_ID:
        HEADING_STYLES.add(UNNUMBERED_ID)
    if UNNUMBERED_HIDDEN_ID:
        HEADING_STYLES.add(UNNUMBERED_HIDDEN_ID)
    HEADING2_STYLES = {HEADING2_ID, "Heading2", "Heading 2", "2"}
    HEADING3_STYLES = {HEADING3_ID, "Heading3", "Heading 3", "3"}
    # digit → 目标一级标题样式 ID（None 表示保持 Heading 1 不变）。
    # 第3节（正文）保留标题1；其他节使用「非编号章节标题」；第5节使用隐藏版本。
    DIGIT_HEADING_STYLE = {
        2: UNNUMBERED_ID,
        3: None,
        4: UNNUMBERED_ID,
        5: UNNUMBERED_HIDDEN_ID,
    }

    # 顺序遍历正文段落；每个章标题与最近的一个 RmdMarker 数字配对。
    # 同时收集所有 marker 段以便最后清理。
    chapters = []         # [(章标题段, 节号), ...]
    markers_to_remove = []
    current_digit = None
    marker_pat = re.compile(r"^RmdMarker(\d+)$")

    for p in list(body.findall(qn("w:p"))):
        sval = _read_para_style(p)
        if sval:
            m = marker_pat.match(sval)
            if m:
                current_digit = int(m.group(1))
                markers_to_remove.append(p)
                continue
        if sval in HEADING_STYLES and current_digit is not None:
            # 按节号→样式映射改写章标题样式。
            target = DIGIT_HEADING_STYLE.get(current_digit)
            # ABSTRACT 标题特殊处理：使用模板中的「ABSTRACT」样式。
            heading_text = _gather_para_text(p).strip()
            if heading_text.upper() == "ABSTRACT" and ABSTRACT_ID:
                target = ABSTRACT_ID
            if target and sval != target:
                ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
                if ps is not None:
                    ps.set(qn("w:val"), target)
            chapters.append((p, current_digit))
        elif current_digit == 4 and sval is not None:
            # 第4节（附录等）的 ##/### 改用 附录标题2/附录标题3。
            ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is not None:
                if sval in HEADING2_STYLES and APPENDIX_H2_ID:
                    ps.set(qn("w:val"), APPENDIX_H2_ID)
                elif sval in HEADING3_STYLES and APPENDIX_H3_ID:
                    ps.set(qn("w:val"), APPENDIX_H3_ID)

    if not chapters:
        for p in markers_to_remove:
            body.remove(p)
        return

    def build_sectpr_for_digit(d):
        idx = max(0, min(d - 1, len(tpl_sectprs) - 1))
        new = copy.deepcopy(tpl_sectprs[idx])
        _rewrite_sectpr_rids(new, tpl_rid_to_target, out_target_to_rid)
        return new

    # 对索引 >= 1 的每个章节，把「上一章」的 sectPr 写入紧邻当前章标题
    # 之前那个段落的 pPr。这样就能正确地结束上一章对应的节。
    marker_set = set(markers_to_remove)
    for i in range(1, len(chapters)):
        p_chap, current_digit = chapters[i]
        prev_digit = chapters[i - 1][1]

        if current_digit == prev_digit:
            # 同一节内：避免插入分节符（不会重启页码），但要插入分页符
            # 让新章另起一页。如果在 TOC（SDT）内则跳过，因为
            # _relocate_toc 已经在前面加过分页符了。
            if p_chap.getparent().tag != qn("w:sdtContent"):
                p_chap.addprevious(_new_page_break_para())
            continue

        prev = p_chap.getprevious()
        while prev is not None and (
            prev.tag != qn("w:p") or prev in marker_set
        ):
            prev = prev.getprevious()
        if prev is None:
            continue
        prev_pPr = _ensure_pPr(prev)
        # 清掉原有的 sectPr（如果有）。
        for old in prev_pPr.findall(qn("w:sectPr")):
            prev_pPr.remove(old)
        prev_pPr.append(build_sectpr_for_digit(prev_digit))

    # 用最后一章的 sectPr 替换 body 级 sectPr。
    last_digit = chapters[-1][1]
    new_body_sect = build_sectpr_for_digit(last_digit)
    for old in body.findall(qn("w:sectPr")):
        body.remove(old)
    body.append(new_body_sect)

    # 清除 marker 段。
    for p in markers_to_remove:
        body.remove(p)

    # 从 styles.xml 里删除 RmdMarker* 样式定义 —— 它们只是后处理过程中
    # 用来标记数字归属的临时样式。
    styles_el = doc.styles.element
    for s in list(styles_el.findall(qn("w:style"))):
        sid = s.get(qn("w:styleId")) or ""
        if re.match(r"^RmdMarker\d+$", sid):
            styles_el.remove(s)
            continue
        name_el = s.find(qn("w:name"))
        nval = name_el.get(qn("w:val")) if name_el is not None else ""
        if nval and re.match(r"^RmdMarker\d+$", nval):
            styles_el.remove(s)


def main() -> None:
    args = [a for a in sys.argv[1:] if a != "--quiet"]
    # --no-bibtex：跳过 BibTeX/citeproc，参考文献按纯文本处理
    use_bibtex = "--no-bibtex" not in args
    args = [a for a in args if a != "--no-bibtex"]
    # --csl=<numeric|note|author-date>；默认 numeric
    csl = "numeric"
    raw = []
    for a in args:
        if a.startswith("--csl="):
            csl = a.split("=", 1)[1]
        elif a in CSL_OPTIONS:
            csl = a
        else:
            raw.append(a)
    if not raw:
        raw = ["docx"]
    pandoc = locate_pandoc()
    formats = [fmt.replace("bookdown::", "").replace("_book", "") for fmt in raw]
    src = fuse_sources(docx_mode=True, use_bibtex=use_bibtex)
    try:
        for f in formats:
            if f == "docx":
                render_docx(pandoc, src, lang="zh", csl=csl, use_bibtex=use_bibtex)
            elif f in ("docx-en", "docx_en"):
                render_docx(pandoc, src, lang="en", csl=csl, use_bibtex=use_bibtex)
            else:
                sys.exit(f"未知输出格式: {f}")
    finally:
        if src.exists():
            try:
                src.unlink()
            except OSError:
                pass


if __name__ == "__main__":
    main()
